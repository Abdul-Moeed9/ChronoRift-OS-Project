from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "Chrono_Rift_Full_Code_Understanding_Guide.docx"


def set_font(run, size=10, bold=False, italic=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size, True)


def subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, 10, False, True)


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, 10)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, 10)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, 10)


def function_block(doc, name, theory, implementation, flow="", demo=""):
    h(doc, name, 4)
    para(doc, "Background concept: " + theory)
    para(doc, "In this code: " + implementation)
    if flow:
        para(doc, "Flow detail: " + flow)
    if demo:
        para(doc, "Demo explanation: " + demo)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def simple_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, head in enumerate(headers):
        table.rows[0].cells[i].text = head
        shade_cell(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    for r_i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, 9, r_i == 0)
    doc.add_paragraph()


doc = Document()
doc.styles["Normal"].font.name = "Arial"
doc.styles["Normal"].font.size = Pt(10)
margins(doc)

title(doc, "Chrono Rift Full Code Understanding Guide")
subtitle(doc, "Theory first, then function-by-function implementation walkthrough")
para(doc, "Group member 1: Abdul Moeed - 24I-0720")
para(doc, "Group member 2: Hamza Sheikh - 24I-0845")
para(doc, "Purpose: This document is a study guide for the project demo. It explains the operating-system concepts first, then connects each concept to the actual code files and functions. Read it in order to understand the flow of the game and the flow of the code.")

h(doc, "1. Big Picture: What This Project Is Doing")
para(doc, "Chrono Rift is a turn-based tactical RPG implemented using operating-system concepts. The game demonstrates process architecture, shared memory, synchronization, multithreading, signals, scheduling, memory management, and deadlock handling.")
para(doc, "The project has three main processes. Arbiter is the parent process and central authority. HIP is the Human Interfacing Process and handles player-side input. ASP is the Automated Strategic Process and handles NPC decisions. The processes communicate through one System V shared memory segment.")
para(doc, "The most important design rule is that Arbiter owns the official game state. HIP and ASP do not directly apply final gameplay effects. They submit a request into shared memory. Arbiter validates that request and applies it.")

h(doc, "2. Core OS Concepts Used")
for head, text in [
    ("Processes", "A process is an independent running program with its own memory space. Arbiter, HIP, and ASP run as separate processes. Arbiter creates HIP and ASP using fork, and the child process uses execl to run the correct executable."),
    ("Shared Memory", "Shared memory is an IPC technique where multiple processes attach to the same memory segment. This project uses shmget and shmat. The shared memory contains one large SharedState structure."),
    ("Semaphores", "Semaphores prevent race conditions and coordinate waiting. GlobalLock protects shared memory. PlayerTurnReady and EnemyTurnReady wake HIP or ASP. ActionReady tells Arbiter that an action request was submitted."),
    ("Threads", "Threads are execution units inside a process. HIP creates one thread per player. ASP creates one thread per enemy. Renderer and deadlock monitor also run as threads inside Arbiter."),
    ("Signals", "Signals are asynchronous notifications. SIGUSR1 is used for stun notification, SIGALRM is used for the ultimate timer, SIGSTOP/SIGCONT pause and resume ASP, and SIGTERM handles quit."),
    ("Scheduling", "Each entity has stamina and speed. Every second, speed is added to stamina. When stamina reaches maximum, the entity receives a ReadyOrder ticket. The smallest ReadyOrder acts first."),
    ("Deadlock", "A deadlock happens when entities wait on each other in a cycle. Here, artifacts are resources. Arbiter builds a wait-for graph and breaks cycles by releasing one artifact."),
    ("Memory Management", "The inventory system models contiguous memory allocation. Weapons require adjacent slots. Fragmentation is handled by swapping weapons to long-term storage.")
]:
    h(doc, "2.x " + head, 2)
    para(doc, text)

h(doc, "3. File and Module Map")
simple_table(doc, ["File / Folder", "Purpose"], [
    ("shared.h", "Defines all shared-memory structures and semaphores."),
    ("..._arbiter/arbiter.cpp", "Main authority: shared memory setup, scheduler, action handling, inventory, artifacts, deadlock, signals, cleanup."),
    ("..._hip/hip.cpp", "Human process: player threads, GUI input request, action submission."),
    ("..._asp/asp.cpp", "Enemy process: enemy threads and AI action submission."),
    ("..._arbiter/renderer.h", "SFML rendering thread and GUI event handling."),
    ("Makefile / Dockerfile", "Build and environment support.")
])

h(doc, "4. shared.h: Shared Memory Layout")
for item in [
    ("WeaponEntry", "A weapon stored in long-term storage needs metadata.", "Stores WeaponId, SlotStart, SlotSize, Damage, and InUse."),
    ("InventorySlot", "The active inventory is a linear slot array.", "Stores WeaponId and OccupiedBy. OccupiedBy points to the starting slot of the weapon."),
    ("LongTermStorage", "Swapped-out weapons need secondary storage.", "Contains Weapons[20] and Count."),
    ("EntityState", "Players and enemies share gameplay attributes.", "Stores HP, stamina, speed, damage, alive/stun fields, pid, weapon/artifact fields, DeathTime, and ReadyOrder."),
    ("ArtifactEntry", "Artifacts are exclusive global resources.", "Stores free/held state, holder pid/entity, player/enemy flag, and whether the artifact is in world."),
    ("ActionRequest", "HIP/ASP need a request format.", "PendingAction stores EntityId, IsPlayer, ActionType, TargetId, WeaponId, LtsIndex, TurnId, and Ready."),
    ("TurnControl", "The game needs to know whose turn is active.", "Stores ActiveEntityId, IsPlayerTurn, Phase, TurnStartTime, and TurnId."),
    ("GuiInput", "Renderer and HIP exchange GUI selections.", "Stores selected action, target, weapon, LTS index, artifact, confirmation flags, and UI phase."),
    ("WeaponDropEvent", "Weapon drops need temporary shared state.", "Stores pending drop, weapon id, target player, and pickup decision."),
    ("GameStatus", "Global game status must be shared.", "Stores kill count, game over flag, winner flag, player/enemy counts, Eclipse state, and SchedulerCounter."),
    ("SharedState", "All shared data is grouped into one segment.", "Contains players, enemies, inventory, artifacts, action/turn/gui state, logs, semaphores, and pids.")
]:
    function_block(doc, item[0], item[1], item[2])

h(doc, "5. Arbiter: Startup and Initialization")
for item in [
    ("main()", "Every program starts from main.", "Seeds randomness with srand(720), asks for party size, randomizes enemy count, creates shared memory, initializes state, forks HIP/ASP, starts renderer/deadlock threads, and enters run_scheduler.", "This is the boot sequence of the entire game.", "Arbiter main builds the game world and launches every other component."),
    ("create_shared_memory()", "Shared memory must exist before other processes attach.", "Calls shmget to create the segment and shmat to attach it. Stores the pointer in g_state."),
    ("init_sync_primitives()", "Semaphores must be initialized before use.", "Initializes GlobalLock, InputLock, PlayerTurnReady, EnemyTurnReady, and ActionReady with pshared = 1."),
    ("init_shared_state()", "Shared memory should start from a clean state.", "Zeroes SharedState, initializes semaphores, resets status, turn, GUI, drop event, and log fields."),
    ("init_players()", "Players follow roll-number stat rules.", "Sets player HP, max stamina 100, speed as 100 / party size, damage 10 for roll 0720, alive/stun/artifact fields, and ReadyOrder."),
    ("init_enemies()", "Enemies follow random stat rules.", "Sets enemy HP, max stamina 150, speed random 10-30, damage 12 for roll 0720, alive/stun/weapon/artifact fields, and ReadyOrder."),
    ("init_inventories()", "Inventories start empty.", "Sets all active slots to WeaponId = -1 and clears LTS entries."),
    ("init_artifacts()", "Artifacts need initial resource states.", "Solar and Lunar start in-world and free. Eclipse starts out of world."),
    ("setup_signals()", "Arbiter needs handlers for quit and timer.", "Registers SIGTERM/SIGINT to sigterm_handler and SIGALRM to sigalrm_handler."),
    ("fork_hip() and fork_asp()", "fork creates child processes and execl runs the child executable.", "Passes shmid as argv so HIP/ASP can attach to the same shared memory."),
    ("print_init_summary()", "Console output helps verify initialization.", "Prints process id, party/enemy counts, and stats.")
]:
    function_block(doc, *item)

h(doc, "6. Arbiter: Cleanup and Signals")
for item in [
    ("cleanup_shared_memory()", "OS resources should be released.", "Destroys semaphores, detaches shared memory, and removes it with shmctl IPC_RMID."),
    ("sigterm_handler()", "Signal handlers should be short.", "Sets g_sigterm_received = 1 so the scheduler can handle quit safely."),
    ("sigalrm_handler()", "SIGALRM is the ultimate timer callback.", "Sends SIGCONT to ASP and clears the frozen flag after alarm(10)."),
    ("wait_for_children()", "Parents should wait for children to avoid zombies.", "Sends SIGCONT/SIGTERM as needed and waits for ASP and HIP using waitpid.")
]:
    function_block(doc, *item)

h(doc, "7. Arbiter: Scheduling Functions")
for item in [
    ("run_scheduler()", "The scheduler is the main game loop.", "Every second it updates stamina, finds a ready entity, starts a turn, wakes HIP/ASP, waits for ActionReady, validates PendingAction, applies the action or timeout-skip, resets turn state, checks game over, and handles drops.", "Player timeout is 30 seconds; enemy timeout is 3 seconds. Short sleeps let Arbiter react early instead of blindly sleeping the whole timeout.", "This function enforces stamina-based serial turn execution."),
    ("refresh_stuns()", "Timed effects need a deadline.", "Gets current time with time(NULL). If StunEndTime <= now, it clears Stunned and gives a new ReadyOrder if stamina is full."),
    ("tick_stamina()", "Stamina fills once per scheduler tick.", "Calls refresh_stuns, adds Speed to stamina for alive non-stunned entities, caps stamina at max, and assigns ReadyOrder when stamina reaches max."),
    ("find_ready_entity(int* outIsPlayer)", "The scheduler must choose the earliest ready entity.", "Scans full-stamina, alive, non-stunned players/enemies and keeps the candidate with the smallest ReadyOrder. outIsPlayer tells whether the returned id is player or enemy.", "bestId means best candidate found so far. -1 means nobody is ready."),
    ("drain_action_ready()", "Old semaphore posts can confuse a new turn.", "Consumes leftover ActionReady posts using sem_trywait and clears PendingAction.Ready."),
    ("pending_action_is_current()", "Requests must match the current turn.", "Checks Ready, EntityId, IsPlayer, and TurnId before Arbiter accepts an action.")
]:
    function_block(doc, *item)

h(doc, "8. Arbiter: Artifact and Deadlock Functions")
for item in [
    ("refresh_entity_artifact_status()", "The code needs a simple summary of artifact ownership.", "Updates HoldsArtifact based on ArtifactHeld[3]."),
    ("entity_has_artifact()", "Many checks need to know if an entity holds one artifact.", "Returns ArtifactHeld[artifactId] for a player or enemy."),
    ("artifact_belongs_to_entity()", "Ownership should be verified from the global artifact table.", "Checks that artifact is held by the exact entity and side."),
    ("try_acquire_artifact()", "Acquiring a resource means checking and updating a resource table.", "If free, marks artifact held and updates entity ArtifactHeld. If held by someone else, records WaitingForArtifact."),
    ("release_artifact()", "Resources must be released on death or deadlock resolution.", "Marks artifact free, clears holder fields, clears entity ownership, removes Solar/Lunar from player inventory/LTS when needed, and logs release."),
    ("clear_artifact_from_lts()", "Released artifacts should not remain in storage.", "Scans LTS and clears matching artifact entries."),
    ("check_and_resolve_deadlock()", "Deadlock detection uses a wait-for graph.", "Combines players/enemies into arrays, creates edges from waiting entities to holders, detects cycles, chooses a victim, releases one artifact, and logs the resolution.", "A cycle like Player -> Enemy -> Player means circular wait."),
    ("deadlock_monitor_func()", "Deadlock checks run in the background.", "Every 2 seconds, locks GlobalLock, calls check_and_resolve_deadlock, unlocks, and repeats until game over."),
    ("try_introduce_eclipse_relic()", "Eclipse is dynamic.", "After at least 3 kills, rolls 30 percent chance to mark Eclipse in-world and free.")
]:
    function_block(doc, *item)

h(doc, "9. Arbiter: Inventory and Memory Management")
for item in [
    ("get_weapon_damage()", "Weapon ids map to fixed damage values.", "Returns damage from the project weapon table."),
    ("get_weapon_slot_size()", "Weapon ids map to fixed contiguous slot sizes.", "Returns slot size from the project weapon table."),
    ("allocate_weapon()", "This is first-fit contiguous allocation.", "Scans from slot 0 for a free block large enough. Marks each slot with WeaponId and OccupiedBy, then returns the start slot."),
    ("free_weapon_at()", "Freeing allocation clears all occupied slots.", "Gets weapon size and clears each slot belonging to the weapon."),
    ("find_free_lts_slot()", "Swap-out needs an empty LTS entry.", "Returns first unused LTS index or -1."),
    ("swap_out_to_lts()", "Swap-out moves active inventory to long-term storage.", "Stores weapon metadata in LTS, increments count, frees active slots, and logs the move."),
    ("find_weapon_start_slots()", "The code must count whole weapons, not individual occupied slots.", "Records unique start slots for weapons in active inventory."),
    ("compact_inventory()", "Compaction moves allocations together.", "Shifts weapons toward the front to reduce gaps. It is a helper for inventory organization."),
    ("auto_swap_out_for_space()", "Fragmentation handling frees the fewest blocking weapons.", "Checks possible windows, counts blocking weapons, chooses the window requiring least swaps, moves those weapons to LTS, and protects Solar/Lunar pairing."),
    ("place_weapon_with_swap()", "Allocation should try normal first, swap only if needed.", "Calls allocate_weapon. If it fails, calls auto_swap_out_for_space and retries."),
    ("swap_in_from_lts()", "Swap In retrieves from long-term storage.", "Validates LTS entry, verifies artifact ownership for Solar/Lunar, places weapon with swap if needed, and clears LTS entry."),
    ("find_first_weapon_in_inventory()", "Use Weapon needs a fallback.", "Returns the first active weapon id or -1 if no weapon exists.")
]:
    function_block(doc, *item)

h(doc, "10. Arbiter: Combat, Actions, Drops, and Game Status")
for item in [
    ("deliver_stun()", "Stun is a timed signal-based status effect.", "Sets Stunned, StunEndTime = time(NULL) + 3, resets ReadyOrder, logs, and sends SIGUSR1 to HIP or ASP."),
    ("roll_weapon_drop()", "Random drops need a weapon id.", "Returns one of the non-artifact weapon ids based on a random roll."),
    ("give_weapon_to_enemy()", "Rejected drops go to an enemy.", "Finds the first alive enemy without a weapon and assigns the dropped weapon."),
    ("handle_enemy_killed()", "Enemy death affects many systems.", "Sets enemy dead, increments kills, releases artifacts, tries Eclipse introduction, and may create a weapon drop."),
    ("handle_player_killed()", "Player death removes them and releases resources.", "Sets player dead and releases any held artifacts."),
    ("process_weapon_drop()", "Drop choice is a temporary interaction.", "Waits for player decision. If accepted, places weapon in inventory. Otherwise gives it to an enemy. Clears DropEvent."),
    ("apply_strike()", "Strike is direct HP damage.", "Player strike damages enemy. Enemy strike damages player and may stun."),
    ("apply_exhaust()", "Exhaust attacks stamina.", "Reduces enemy stamina by player damage and resets player stamina."),
    ("apply_use_weapon()", "Use Weapon applies weapon damage.", "Finds requested/first weapon, verifies artifact ownership, applies weapon damage, may stun enemy, and handles death."),
    ("apply_swap_in()", "Swap In costs a full turn.", "Calls swap_in_from_lts and resets player stamina."),
    ("apply_heal()", "Heal restores 10 percent HP.", "Caps healing at MaxHp and resets stamina."),
    ("apply_skip()", "Skip keeps half stamina.", "Sets stamina to MaxStamina / 2 and clears ReadyOrder."),
    ("apply_ultimate()", "Ultimate requires Solar and Lunar.", "Verifies both artifacts in inventory and ownership table, sends SIGSTOP to ASP, starts alarm(10), and resets stamina."),
    ("apply_acquire_artifact()", "Artifact acquisition is a turn action.", "Calls try_acquire_artifact and places Solar/Lunar as weapons if newly acquired by player."),
    ("apply_action()", "This dispatches PendingAction to the correct handler.", "Reads action fields, validates target/alive/stunned state, calls the matching action function, then clears PendingAction.Ready."),
    ("check_game_over()", "The game ends on win or loss.", "Win at 10 enemy kills. Lose if all players die."),
    ("respawn_dead_enemies()", "The enemy pool continues until 10 kills.", "After a delay, revives dead enemy slots with new stats if the game is not over.")
]:
    function_block(doc, *item)

h(doc, "11. HIP Process: Function-by-Function")
para(doc, "HIP handles player-side turn processing. It waits for Arbiter to announce a player turn, wakes the correct player thread, gets GUI input, and submits PendingAction.")
for item in [
    ("HIP main()", "HIP is a child process that attaches to shared memory.", "Reads shmid from argv, attaches with shmat, installs signals, creates dispatch thread and player threads, joins them, destroys local semaphores, and detaches shared memory."),
    ("sigterm_handler()", "HIP must stop on quit.", "Sets g_running = 0."),
    ("sigusr1_handler()", "HIP receives stun notification.", "Sleeps 3 seconds and clears expired player stun flags."),
    ("submit_action()", "Player choices become shared-memory requests.", "Validates current turn, fills PendingAction, sets Ready, unlocks, and posts ActionReady."),
    ("request_gui_action_menu()", "HIP asks UI for action input.", "Sets Gui.Phase = 1 and clears old GUI selections."),
    ("request_gui_drop_decision()", "HIP asks UI for drop decision.", "Sets Gui.Phase = 6 and clears drop confirmation."),
    ("wait_for_gui_action()", "HIP waits for renderer confirmation.", "Loops until Gui.Confirmed or game end."),
    ("wait_for_gui_drop()", "HIP waits for drop accept/reject.", "Loops until Gui.DropConfirmed or game end."),
    ("clear_gui_state()", "GUI state resets after use.", "Sets phase and confirmations back to idle."),
    ("process_turn()", "This converts GUI selection into an action request.", "Requests menu, waits, reads chosen fields, maps them to submit_action, or sends SIGTERM for Quit."),
    ("check_weapon_drop()", "After a kill, player may need to choose a drop.", "If DropEvent is pending for this player, requests drop decision and writes result."),
    ("player_thread_func()", "Each player thread waits for its own turn.", "Waits on local player semaphore, verifies active turn, processes turn, then waits again."),
    ("player_dispatch_func()", "Routes global player turn to exact thread.", "Waits on PlayerTurnReady, reads ActiveEntityId, and posts g_player_turn[activeId].")
]:
    function_block(doc, *item)

h(doc, "12. ASP Process: Function-by-Function")
para(doc, "ASP handles automated enemy decisions. Like HIP, it submits action requests rather than applying official effects.")
for item in [
    ("ASP main()", "ASP is a child process that attaches to shared memory.", "Reads shmid, attaches, installs signals, seeds randomness, creates dispatch thread and enemy threads, joins them, destroys semaphores, and detaches."),
    ("sigterm_handler()", "ASP stops on quit.", "Sets g_running = 0."),
    ("sigusr1_handler()", "ASP receives stun notification.", "Sleeps 3 seconds and clears expired enemy stun flags."),
    ("pick_alive_player()", "Enemy Strike needs a living target.", "Builds a list of alive players and randomly chooses one."),
    ("submit_enemy_action()", "Enemy choices become PendingAction requests.", "Validates current enemy turn, fills PendingAction, and posts ActionReady."),
    ("enemy_do_strike()", "Enemy strike submits attack.", "Chooses an alive player and submits action type 0, or skip if no target."),
    ("enemy_do_skip()", "Enemy skip submits action type 5.", "Submits skip for the enemy."),
    ("enemy_do_acquire_artifact()", "Enemy may acquire free artifacts.", "Finds a free in-world artifact and submits acquire; otherwise strikes."),
    ("enemy_decide_action()", "This is enemy AI.", "Sleeps 0-2 seconds, then randomly chooses strike, skip, or acquire artifact."),
    ("enemy_thread_func()", "Each enemy thread waits for its own turn.", "Waits on local enemy semaphore, verifies active enemy id, then calls enemy_decide_action."),
    ("enemy_dispatch_func()", "Routes global enemy turn to exact thread.", "Waits on EnemyTurnReady, reads ActiveEntityId, and posts g_enemy_turn[activeId]."),
    ("attach_shared_memory()", "ASP attaches to the shared state.", "Calls shmat and stores pointer in g_state."),
    ("setup_signals()", "ASP installs handlers.", "Registers SIGTERM and SIGUSR1."),
    ("create_enemy_threads() / join_enemy_threads()", "Thread management helpers.", "Creates one pthread per enemy and joins each one before exit."),
    ("detach_shared_memory()", "ASP detaches before exit.", "Calls shmdt if attached.")
]:
    function_block(doc, *item)

h(doc, "13. Renderer: Function Groups")
para(doc, "Renderer is large because it contains many drawing helpers. Understand its role rather than memorizing every coordinate. It runs as a pthread inside Arbiter, snapshots shared memory, draws the UI, and records GUI choices.")
for item in [
    ("render_thread_func()", "Rendering should not block the scheduler.", "Receives SharedState*, creates SFML window, handles events, copies SharedState under GlobalLock into snapshot, draws from snapshot, and exits on game over/window close.", "Snapshotting keeps lock time short."),
    ("render_frame()", "A frame is one complete screen draw.", "Calls draw functions for top bar, players, enemies, artifacts, logs, menu, drop notification, inventory overlay, and game-over overlay."),
    ("draw_top_bar()", "Global state should be visible.", "Shows title, kill count, active turn, victory/defeat, and Eclipse status."),
    ("draw_player_battleground() / draw_enemy_battleground()", "Entities need visual HP/stamina state.", "Draws sprites, bars, active highlighting, dead/stun labels, and enemy selectable regions."),
    ("draw_action_menu()", "Player actions need GUI controls.", "Draws Strike, Exhaust, Use Weapon, Swap In, Heal, Skip, Ultimate, Acquire Artifact, and Quit controls."),
    ("draw_inventory_overlay()", "Inventory visualization supports memory-management demo.", "Shows 20 active slots and LTS entries."),
    ("draw_log_panel()", "Logs explain runtime behavior.", "Displays recent ActionLog entries."),
    ("draw_weapon_drop_notification()", "Drop choices need UI.", "Shows dropped weapon and accept/reject options."),
    ("draw_game_over_overlay()", "End state should be obvious.", "Shows victory or defeat overlay."),
    ("finalize_gui_action()", "Renderer writes chosen action to shared GUI state.", "Stores action/target/weapon/LTS/artifact selections and sets Confirmed = 1."),
    ("finalize_drop_choice()", "Renderer writes drop choice.", "Stores DropChoice and sets DropConfirmed = 1."),
    ("handle_action_chosen() and related handlers", "Events must become game selections.", "Advance local GUI phase and eventually call finalize_gui_action."),
    ("handle_mouse_click(), handle_mouse_move(), handle_key_press()", "SFML input must map to menu state.", "Route mouse/keyboard events according to current GUI phase."),
    ("update_gui_phase()", "Renderer local phase follows shared Gui.Phase.", "Keeps displayed UI synced with what HIP is waiting for."),
    ("stop_renderer() / is_renderer_running()", "Arbiter needs control over renderer lifetime.", "Set or return the render running flag.")
]:
    function_block(doc, *item)

h(doc, "14. Full Game Flow From Start to End")
for step in [
    "User runs ./arbiter.out.",
    "Arbiter seeds randomness with roll number 720 and asks for party size.",
    "Arbiter creates shared memory and initializes SharedState.",
    "Arbiter forks HIP and ASP, passing shmid to both.",
    "HIP and ASP attach to shared memory and create entity threads.",
    "Arbiter starts renderer thread and deadlock monitor thread.",
    "Arbiter enters run_scheduler.",
    "Every second, tick_stamina updates stamina and assigns ReadyOrder tickets.",
    "find_ready_entity selects the full-stamina entity with the smallest ReadyOrder.",
    "Arbiter sets Turn fields and wakes HIP or ASP using a semaphore.",
    "HIP/ASP dispatches the turn to the correct player/enemy thread.",
    "The active thread submits PendingAction and posts ActionReady.",
    "Arbiter validates PendingAction with TurnId and applies it through apply_action.",
    "Action handlers update HP, stamina, inventory, artifacts, logs, drops, and game status.",
    "Renderer continuously snapshots SharedState and displays the current state.",
    "Deadlock monitor periodically checks artifact circular wait.",
    "Game ends when 10 enemies are killed, all players die, or quit occurs.",
    "Arbiter signals children, joins threads/processes, detaches and removes shared memory."
]:
    numbered(doc, step)

h(doc, "15. Rubric Mapping Summary")
simple_table(doc, ["Rubric Area", "Where It Appears", "Demo Defense"], [
    ("Process Architecture & IPC", "Arbiter main, fork_hip, fork_asp, SharedState", "Three separate processes communicate using System V shared memory."),
    ("Synchronization", "GlobalLock, ActionReady, PlayerTurnReady, EnemyTurnReady", "Semaphores protect shared memory and coordinate turn/action flow."),
    ("Multithreading", "HIP player threads, ASP enemy threads, renderer/deadlock threads", "Each player/enemy has its own pthread; dispatchers wake only the active one."),
    ("Scheduling", "run_scheduler, tick_stamina, find_ready_entity", "Stamina fills per second; ReadyOrder selects earliest ready entity."),
    ("Signals", "deliver_stun, apply_ultimate, signal handlers", "SIGUSR1 handles stun notification, SIGALRM manages ultimate timer, SIGTERM handles quit."),
    ("Deadlock", "Artifacts table, check_and_resolve_deadlock", "Wait-for graph detects artifact circular wait and releases one resource."),
    ("Memory Management", "allocate_weapon, auto_swap_out_for_space, swap_in_from_lts", "20-slot inventory uses first-fit, fragmentation handling, and LTS."),
    ("Gameplay", "apply_action and handlers", "HIP/ASP request actions; Arbiter validates and applies effects."),
    ("Rendering", "render_thread_func and draw functions", "Dedicated SFML thread snapshots shared state and displays it."),
    ("Environment", "Dockerfile and Makefile", "Build produces arbiter.out, hip.out, asp.out.")
])

h(doc, "16. Common Demo Questions")
for q, a in [
    ("What is SharedState?", "The complete shared memory structure containing entities, inventory, artifacts, action request, turn state, GUI state, logs, semaphores, and pids."),
    ("What is PendingAction?", "The action request written by HIP or ASP and read by Arbiter."),
    ("What is ActionReady?", "The semaphore notification that tells Arbiter PendingAction has been submitted."),
    ("What is TurnId?", "A unique number per turn that prevents late/stale actions from being accepted."),
    ("What is Turn.Phase?", "Phase 1 means active turn; Phase 0 means no active turn."),
    ("Why not sleep(3) for ASP?", "Short sleeps let Arbiter react as soon as action arrives while still enforcing a 3-second maximum."),
    ("What is ReadyOrder?", "A scheduling ticket assigned when stamina becomes full. Smallest ticket acts first."),
    ("Is ReadyOrder Lamport Bakery?", "No. It is a turn-scheduling ticket mechanism, not mutual exclusion."),
    ("How does stun expire?", "StunEndTime is current time plus 3. refresh_stuns clears stun once current time reaches that value."),
    ("How is deadlock detected?", "The code builds a wait-for graph from WaitingForArtifact and artifact holders. A cycle means deadlock."),
    ("Why is renderer.h included?", "arbiter.cpp includes renderer.h, so render_thread_func is visible when compiling Arbiter."),
    ("Who owns game logic?", "Arbiter. HIP and ASP only submit requests.")
]:
    para(doc, "Q: " + q)
    para(doc, "A: " + a)

h(doc, "17. What To Study First If Time Is Short")
for item in [
    "First: shared.h, especially SharedState, EntityState, ActionRequest, TurnControl, ArtifactEntry.",
    "Second: Arbiter main and run_scheduler.",
    "Third: tick_stamina, find_ready_entity, SchedulerCounter, ReadyOrder.",
    "Fourth: HIP process_turn/submit_action and ASP enemy_decide_action/submit_enemy_action.",
    "Fifth: apply_action and action handlers.",
    "Sixth: inventory and deadlock functions.",
    "Last: renderer.h surface flow, especially render_thread_func and snapshot logic."
]:
    bullet(doc, item)

h(doc, "18. Safe Demo Language")
para(doc, "Say that Arbiter is the authority and HIP/ASP submit requests. Say that shared memory stores the state and semaphores coordinate access. Say that ReadyOrder is a scheduling ticket mechanism. Say that renderer snapshots state and displays it, but does not apply game rules. Avoid claiming bonus multiplayer unless it is implemented.")

margins(doc)
doc.save(OUT)
print(OUT)
